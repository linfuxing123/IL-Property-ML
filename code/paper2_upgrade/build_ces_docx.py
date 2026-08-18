#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_ces_docx.py — 从 manuscript_ces.md 构建 Chemical Engineering Science 投稿版 DOCX

版式：Times New Roman 12pt / 双倍行距 / 单栏 / 标题加粗 / markdown 表格转 Word 表格 /
图件（Fig.1-6）嵌入文末并附图题。
用法: python build_ces_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
SRC = BASE / "manuscript_ces.md"
OUT = BASE / "manuscript_ces.docx"
FIGS = BASE / "results"
BLACK = RGBColor(0, 0, 0)
FIG_FILES = ["fig_scaling_laws.png", "fig_leakage_tax.png", "fig_cold_start.png",
             "fig_acquisition_v2.png", "fig_coverage_hist.png", "fig_leaderboard.png"]


def set_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def add_para(doc, text, size=12, bold=False, italic=False, align=None, double=True,
             before=0, after=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 2.0 if double else 1.15
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); set_run(r, size, True, False)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); set_run(r, size, False, True)
        else:
            r = p.add_run(part); set_run(r, size, bold, italic)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ""
        r = c.paragraphs[0].add_run(h); set_run(r, 10.5, True)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            r = c.paragraphs[0].add_run(str(v)); set_run(r, 10.5)
    doc.add_paragraph()


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)

    in_table = []
    pending_caption = None
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        i += 1
        if not line.strip():
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if re.fullmatch(r":?-+:?", cells[0]) or all(re.fullmatch(r":?-+:?", c) for c in cells):
                continue  # 分隔行
            if not in_table:
                in_table = [cells]
            else:
                in_table.append(cells)
            continue
        else:
            if in_table:
                add_table(doc, in_table[0], in_table[1:])
                in_table = []
        if line.startswith("#"):
            title = line.lstrip("# ").strip()
            if line.startswith("## "):
                add_para(doc, title, 14, True, False, before=10, after=4)
            else:
                add_para(doc, title, 16, True, False, before=12, after=6)
            continue
        if line.startswith("**Table"):
            add_para(doc, line, 11, True, False, double=False, after=2)
            continue
        if line.startswith("\\*"):
            add_para(doc, line.lstrip("\\*"), 10.5, False, True, double=False)
            continue
        if line.startswith("**Fig."):
            # 图题: 在文末统一放图, 这里保留图题文字并随后插图像
            caption = line.strip("*")
            add_para(doc, caption, 11, False, True, double=False, before=8, after=2)
            # 找对应图文件: Fig.1 -> fig_scaling_laws.png 按顺序
            m = re.match(r"Fig\. (\d)", caption)
            if m:
                k = int(m.group(1)) - 1
                if 0 <= k < len(FIG_FILES):
                    fpath = FIGS / FIG_FILES[k]
                    if fpath.exists():
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        r.add_picture(str(fpath), width=Inches(5.8))
            continue
        # 普通段落（合并行内换行）
        para = line
        while i < n and lines[i].strip() and not lines[i].startswith(("#", "|", "**Table", "**Fig.")):
            para += " " + lines[i].strip()
            i += 1
        add_para(doc, para, 12)

    if in_table:
        add_table(doc, in_table[0], in_table[1:])

    OUT.write_bytes(b"")
    doc.save(OUT)
    npar = len(doc.paragraphs)
    print(f"saved {OUT} | paragraphs {npar} | tables {len(doc.tables)}")


if __name__ == "__main__":
    main()
