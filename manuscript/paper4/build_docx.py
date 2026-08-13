#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_docx.py - 从 science_manuscript.md 构建 Science 投稿版 DOCX

样式 token (学术稿件 preset, named override):
  page: US Letter, 1in margins
  base: Times New Roman 12pt, double spacing
  headings: TNR bold black (H1 16pt / H2 14pt / H3 12pt)
  tables: TNR 10.5pt, single grid, header bold
  figures: centered, 6.0in width
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "science_manuscript.md"
OUT = ROOT / "science_manuscript_4.docx"
FIGS = ROOT / "figures"

BLACK = RGBColor(0, 0, 0)


def set_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def add_para(doc, text, size=12, bold=False, italic=False, align=None,
             double=True, before=0, after=0, keep_lines=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if double:
        pf.line_spacing = 2.0
    else:
        pf.line_spacing = 1.15
    if keep_lines:
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        kl = OxmlElement("w:keepLines")
        pPr.append(kl)
    # 处理内联 **bold** 和 *italic*
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_run(r, size, bold or True)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            set_run(r, size, bold, italic or True)
        else:
            r = p.add_run(part)
            set_run(r, size, bold, italic)
    return p


def parse_table(lines):
    """解析 markdown 表格行 -> 单元格列表"""
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln.startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c.strip()):
            continue  # 分隔行
        rows.append(cells)
    return rows


def build():
    doc = Document()
    # 页面设置: Letter 1in
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1.0))
    # Normal 样式
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.paragraph_format.line_spacing = 2.0
    st.paragraph_format.space_after = Pt(0)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    in_table = False
    table_buf = []
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if s.startswith("|"):
            table_buf.append(ln)
            i += 1
            continue
        if table_buf:
            rows = parse_table(table_buf)
            table_buf = []
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=ncol)
                t.style = "Table Grid"
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.autofit = True
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        c = t.rows[ri].cells[ci]
                        c.paragraphs[0].text = ""
                        r = c.paragraphs[0].add_run(cell.replace("**", ""))
                        set_run(r, 10.5, bold=(ri == 0))
                        c.paragraphs[0].paragraph_format.line_spacing = 1.15
                        c.paragraphs[0].paragraph_format.space_before = Pt(2)
                        c.paragraphs[0].paragraph_format.space_after = Pt(2)
                add_para(doc, "", double=False, after=6)
        if not s:
            i += 1
            continue
        if s.startswith("![Figure"):
            m = re.search(r"\((figures/[^)]+)\)", s)
            if m:
                img = FIGS / Path(m.group(1)).name
                if img.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(str(img), width=Inches(5.8))
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if s.startswith("# "):
            add_para(doc, s[2:], size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     double=True, before=0, after=6)
        elif s.startswith("**One-sentence summary:**"):
            add_para(doc, s, size=12, double=False, before=4, after=4)
        elif s.startswith("**Authors:**") or s.startswith("**Affiliation:**") or s.startswith("**Corresponding author:**"):
            add_para(doc, s, size=12, double=False)
        elif s.startswith("## "):
            add_para(doc, s[3:], size=14, bold=True, double=True, before=12, after=4)
        elif s.startswith("### "):
            add_para(doc, s[4:], size=12, bold=True, double=True, before=8, after=2)
        elif s.startswith("**") and s.endswith("**") and len(s) < 200:
            add_para(doc, s, size=12, bold=True, double=False, before=6, after=2)
        else:
            add_para(doc, s)
        i += 1
    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    build()

